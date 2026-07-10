#!/usr/bin/env python3
"""生成 v2.2.0 完整版本文档 .docx"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# 使用系统默认字体，不指定特定字体名
style = doc.styles['Normal']
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)

def p(text):
    doc.add_paragraph(text)

def code(text):
    r = doc.add_paragraph().add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(8.5)

def make_table(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers), style='Light Grid Accent 1')
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i+1].cells[j].text = str(val)
    return t

# ============ 封面 ============
for _ in range(6):
    doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Devops-Glue v2.2.0'); r.font.size = Pt(28); r.font.bold = True
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('版本升级说明'); r.font.size = Pt(18)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('\n\n基准: v2.1.2  |  2026-07-10  |  25 文件  +1692/-459'); r.font.size = Pt(10)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('ProviderRegistry · Gitea · 硬编码清理 · 日志修复 · 首页重构'); r.font.size = Pt(10)
doc.add_page_break()

# ============ 目录 ============
doc.add_heading('目录', level=1)
for item in ['一、版本概述', '二、核心架构变更', '三、新增功能', '四、问题修复',
             '五、job_git_map 字段规范', '六、环境变量完整清单', '七、对外接口变化',
             '八、版本兼容性', '九、测试结果', '十、变更文件索引']:
    doc.add_paragraph(item)
doc.add_page_break()

# ============ 一 ============
doc.add_heading('一、版本概述', level=1)
p('v2.2.0 是一次架构升级版本，核心目标是将 Git 平台层从"硬编码字符串匹配"改造为"可扩展注册表模式"，同时清理历史遗留的硬编码、空异常吞噬等问题。变更规模：25 个文件，+1692 行，-459 行。新增 2 个类文件（GiteaService、ProviderRegistry），重构 15 个已有文件。')
p('')
p('核心亮点：')
p('  - ProviderRegistry 注册表：平台检测从硬编码 switch 升级为可插拔注册')
p('  - Gitea 适配器：自建 Git 服务完整支持，与 GitLab 同级')
p('  - 自定义平台：实现接口 + settings 配置即可接入，不改源码')
p('  - 硬编码清零：移除 5 处 IP/占位符，全部由配置驱动')
p('  - 空 catch 清零：9 处沉默异常全部补 Logger')
p('  - 首页重构：树形展示各服务状态+版本号，仅显示已配置平台')
doc.add_page_break()

# ============ 二 ============
doc.add_heading('二、核心架构变更', level=1)

doc.add_heading('2.1 ProviderRegistry 注册表', level=2)
p('【问题】MapService::detectPlatform() 硬编码字符串匹配关键词，兜底 return "gitlab"。自建 GitLab/Gitea 的 URL 不含平台关键词时被静默误判；新增平台必须改 MapService 源码。')
p('【方案】新建 ProviderRegistry（src/Service/Git/ProviderRegistry.php，105 行）：')
code('  register(name, matcher, factory) - 注册平台（内置 + 自定义均可）')
code('  detect(url) - 遍历 matcher，返回平台名 / 抛明确 ApiException')
code('  create(name) - 调用 factory 闭包，返回 GitProviderInterface')
code('  getRegisteredNames() - 返回已注册平台列表')
p('ProviderRegistry 只注册已配置的平台（isPlatformConfigured 检查），未配不注册、不检测、不显示。')

doc.add_heading('2.2 GitProviderInterface 接口扩展', level=2)
make_table(['方法', '返回值', '说明'], [
    ['getBranches(repository)', 'string[]', '原有，获取分支列表'],
    ['getName()', 'string', '新增，平台名: gitlab|gitee|github|gitea'],
    ['matchUrl(url)', 'bool', '新增，判断 URL 是否属于该平台'],
    ['getApiVersion()', 'string', '新增，API 版本: v4|v5|v3|v1'],
])

doc.add_heading('2.3 平台检测决策树', level=2)
p('(1) URL 含 gitlab/gitee/github/gitea 关键词?')
p('    Yes -> ProviderRegistry.detect() 精确匹配, platform_source="auto", detection_method="exact"')
p('    No  -> 日志 warn, 回退 DEFAULT_GIT_PLATFORM, detection_method="fallback"')
p('(2) job_git_map 中手动指定 git_platform? -> 覆盖自动检测, platform_source="manual"')
p('')
p('错误消息 3 级区分：')
make_table(['级别', '触发条件', '消息内容'], [
    ['L1', '手动指定 git_platform 但未配置', '"Job xxx 指定了 git_platform=\'gitea\'，但该平台尚未配置。请在 .env 中设置 GITEA_BASE_URL..."'],
    ['L2', '自动 fallback 后 API 失败', '"gitlab 分支查询失败。提示：此平台由 DEFAULT_GIT_PLATFORM 兜底识别...请修改"'],
    ['L3', '精确匹配后 API 失败', '原始错误，不追加兜底提示'],
])

doc.add_heading('2.4 可扩展设计', level=2)
p('内置平台（gitlab/gitee/github/gitea）：通过 .env 环境变量配置，在 container.php 中注册。自定义平台：通过 settings.php 的 custom_providers 数组注册，只需实现 GitProviderInterface 接口。两者共用同一个 ProviderRegistry::register()，仅配置来源不同。')
doc.add_page_break()

# ============ 三 ============
doc.add_heading('三、新增功能', level=1)

doc.add_heading('3.1 Gitea 平台支持', level=2)
p('Gitea 是广泛使用的自建 Git 服务。v2.2.0 新增 GiteaService 适配器，与 GitLab 同级。')
make_table(['', 'GitLab 自建', 'Gitea 自建', 'GitHub SaaS', 'Gitee SaaS'], [
    ['分支接口', '/api/v4/projects/{id}/.../branches', '/api/v1/repos/{owner}/{repo}/branches', '/repos/{owner}/{repo}/branches', '/repos/{owner}/{repo}/branches'],
    ['认证头', 'PRIVATE-TOKEN', 'Authorization: token', 'Authorization: token', 'Authorization: token'],
    ['URL 检测', '不可靠', '不可靠', '可靠 (github.com)', '可靠 (gitee.com)'],
    ['API 版本', 'v4', 'v1', 'v3', 'v5'],
])
p('重要：自建平台 URL 检测不可靠，必须靠 DEFAULT_GIT_PLATFORM 或 job_git_map.git_platform 兜底。')

doc.add_heading('3.2 服务版本号显示', level=2)
p('health 接口新增三个字段：jenkins_version（自动探测 X-Jenkins header）、harbor_version（v1/v2）、git[].api_version。')
p('首页树形效果：Jenkins v2.555.3 | Git(gitlab v4) | Git(gitee v5) | Git(github v3) | Harbor v2')

doc.add_heading('3.3 自定义 Git 平台接入', level=2)
p('无需修改源码：(1) 实现 GitProviderInterface 写 Provider 类 (2) settings.php 的 custom_providers 注册。凭证通过 env() 从 .env 读取，不写入 Git。')

doc.add_heading('3.4 首页重构', level=2)
p('树形展示 + 版本号 + 仅已配置平台显示 + 1080p 笔记本一屏。')
doc.add_page_break()

# ============ 四 ============
doc.add_heading('四、问题修复', level=1)

doc.add_heading('4.1 硬编码清理', level=2)
make_table(['位置', '原来', '现在'], [
    ['MapService detectPlatform()', '匹配 "192.168.137.5:8082"', '移除'],
    ['MapService 兜底', 'return "gitlab"', 'ProviderRegistry.detect()'],
    ['container.php Harbor', 'base_uri: http://192.168.137.5', '""'],
    ['settings.php Jenkins', 'url: http://Jenkins_URL', '""'],
    ['settings.php Harbor', 'url: http://Harbor_URL', '""; 删除未使用的 repo 字段'],
    ['测试脚本', 'baseUrl 硬编码', '$argv[1] / 环境变量'],
])

doc.add_heading('4.2 空 catch + Logger', level=2)
p('JenkinsService 4 处 + HarborService 5 处空 catch 全部补 Logger。所有 Service（Jenkins/Harbor/Git/Map）通过 setLogger() 注入，每个 catch 记录 debug/warning 含 URL、方法名、异常消息。')

doc.add_heading('4.3 GitHub API 版本', level=2)
p('显示: api_version="v3" | HTTP: X-GitHub-Api-Version:2022-11-28 | URL: 不拼 /api/v3 | 分页: MAX_PAGES=20')

doc.add_heading('4.4 CSRF 403 区分', level=2)
p('拿到 crumb 后仍 403 -> "权限不足或 Job 配置错误" | 无法获取 crumb -> "权限不足，且无法获取 CSRF crumb"')

doc.add_heading('4.5 CI PHP 7.4 -> 8.3', level=2)
p('.github/workflows/security.yml 3 处对齐。代码使用 match/str_contains 等 PHP 8.0 语法，7.4 lint 无意义。')

doc.add_heading('4.6 Docker 时区', level=2)
p('docker-compose.yml: TZ=Asia/Shanghai')
doc.add_page_break()

# ============ 五 ============
doc.add_heading('五、job_git_map 字段规范', level=1)
make_table(['字段', '必填', '说明'], [
    ['job_name', 'Yes', 'Jenkins Job 完整路径，唯一标识'],
    ['git_platform', '', '自建实例强烈建议。SaaS 可省略。可选: gitlab|gitee|github|gitea'],
    ['git_remote', '', '不填从 Jenkins SCM 自动获取'],
    ['project_id', '', 'GitLab: 不填自动 API 查'],
    ['web_url', '', '项目主页链接，仅展示'],
    ['current_path', '', '不填从 git_remote 推导'],
    ['harbor_repository', '', '关联 Harbor 仓库 "project/repository"，仅展示'],
    ['api_version', '', '纯元数据，不影响 API 路由'],
])
doc.add_page_break()

# ============ 六 ============
doc.add_heading('六、环境变量完整清单', level=1)
make_table(['变量', '默认值', '必填', '说明'], [
    ['JENKINS_BASE_URL', '', 'Yes', 'Jenkins 地址'],
    ['JENKINS_USER', 'admin', '', ''],
    ['JENKINS_TOKEN', '', 'Yes', 'Jenkins API Token'],
    ['GITLAB_BASE_URL', '', 'Yes', '自建 GitLab 地址'],
    ['GITLAB_TOKEN', '', 'Yes', 'GitLab Access Token'],
    ['GITEE_BASE_URL', 'https://gitee.com/api/v5', '', 'SaaS 默认可用'],
    ['GITEE_TOKEN', '', '', ''],
    ['GITHUB_BASE_URL', 'https://api.github.com', '', 'SaaS 默认可用'],
    ['GITHUB_TOKEN', '', '', ''],
    ['GITEA_BASE_URL', '', '', 'v2.2.0 新增; 自建 Gitea'],
    ['GITEA_TOKEN', '', '', 'v2.2.0 新增'],
    ['DEFAULT_GIT_PLATFORM', 'gitlab', '', 'v2.2.0 新增; URL 无法识别时兜底'],
    ['HARBOR_BASE_URL', '', 'Yes', 'Harbor 地址'],
    ['HARBOR_USER', 'admin', '', ''],
    ['HARBOR_PASSWORD', '', '', ''],
    ['APP_ENV', 'production', '', ''],
    ['BUILD_TIMEOUT', '300', '', ''],
    ['LOG_PATH', '/data/logs/ci-platform/', '', ''],
])
doc.add_page_break()

# ============ 七 ============
doc.add_heading('七、对外接口变化', level=1)
make_table(['接口', '变化'], [
    ['GET /api/health', '新增 jenkins_version, harbor_version; git[] 新增 api_version; 未配不出现'],
    ['GET /api/main/git/platforms', '含 gitea+自定义; GitHub URL 不拼版本路径'],
    ['GET /api/git/{job}/branches', '错误 3 级区分 (手动未配/兜底失败/精确失败)'],
    ['首页', '树形+版本号+仅已配置'],
])
doc.add_page_break()

# ============ 八 ============
doc.add_heading('八、版本兼容性', level=1)
make_table(['服务', '文档声明', '实际下限', '上限', '说明'], [
    ['Jenkins', 'v2.555.3+', '2.60', 'latest', '/api/json 稳定; 4 层降级'],
    ['GitLab', 'v17.0', '9.0', '17.x', 'API v4 无破坏变更; 无 v5 计划'],
    ['Gitee', 'API v5', 'API v5', '当前', 'SaaS 稳定'],
    ['GitHub', 'API v3', 'API v3', '当前', 'HTTP:2022-11-28, 显示:v3'],
    ['Gitea', 'API v1', 'API v1', '当前', 'v2.2.0 新增'],
    ['Harbor', 'v1.10.1/v2.x', '1.10.0/2.0.0', '2.x', '自动探测 v1/v2'],
])
doc.add_page_break()

# ============ 九 ============
doc.add_heading('九、测试结果', level=1)
p('环境: localhost:8080 | Jenkins 2.555.3 | Harbor v2 | GitLab + Gitee | 2026-07-09')
make_table(['模块', '测试', '通过', '备注'], [
    ['Infra', '4', '4', 'health(含版本号) + swagger + openapi + CORS'],
    ['Main', '5', '5', 'Job列表(3格式) + 映射 + 平台列表 + 发现'],
    ['Jenkins', '27', '27', '3 Job x 9 接口'],
    ['Git', '3', '2', '1 个预期 404 (gitea 未配，明确报错)'],
    ['Harbor', '5', '5', '项目+仓库+Tag+扫描+报告'],
    ['总计', '44', '43', '通过率 97.7%; 1 个为设计行为非 bug'],
])
p('关键验证: jenkins_version 2.555.3 | harbor_version v2 | 未配 Gitea 不出现 | 手动指定未配平台 -> 明确错误消息')
doc.add_page_break()

# ============ 十 ============
doc.add_heading('十、变更文件索引', level=1)
make_table(['Op', '文件', '说明'], [
    ['M', 'config/.env.example', '+GITEA_BASE_URL, GITEA_TOKEN, DEFAULT_GIT_PLATFORM'],
    ['M', 'config/AppConfig.php', '+getGiteaConfig, getDefaultGitPlatform, GitHub no-path-fix'],
    ['M', 'config/container.php', 'ProviderRegistry + 条件注册 + Logger 注入 + Harbor IP 清理'],
    ['M', 'config/settings.php', '+gitea +default_platform +custom_providers; bogus 清理; 字段表'],
    ['M', 'config/settings.example.php', '同步'],
    ['M', 'docker-compose.yml', 'TZ=Asia/Shanghai'],
    ['M', 'readme.md', 'v2.2.0; Gitea; 自定义平台; 字段表'],
    ['M', '.github/workflows/security.yml', 'PHP 7.4->8.3'],
    ['M', 'src/Service/Git/GitProviderInterface.php', '+getName, matchUrl, getApiVersion'],
    ['M', 'src/Service/Git/GitlabService.php', '新接口 + Logger'],
    ['M', 'src/Service/Git/GiteeService.php', '新接口 + Logger'],
    ['M', 'src/Service/Git/GithubService.php', '新接口 + MAX_PAGES + header + Logger'],
    ['A', 'src/Service/Git/GiteaService.php', '新建, 70 行'],
    ['A', 'src/Service/Git/ProviderRegistry.php', '新建, 105 行'],
    ['M', 'src/Service/Git/GitProviderFactory.php', '委托 Registry'],
    ['M', 'src/Service/MapService.php', 'Registry + 移除硬编码 + default_platform'],
    ['M', 'src/Service/GitService.php', 'Registry + 错误 3 级区分'],
    ['M', 'src/Service/JenkinsService.php', 'Logger + 4 catch + CSRF + getVersion()'],
    ['M', 'src/Service/HarborService.php', 'Logger + 5 catch + getApiVersion()'],
    ['M', 'src/Controller/MainController.php', 'health: 过滤 + 版本号'],
    ['M', 'templates/openapi.json', 'v2.2 + Gitea'],
    ['M', 'templates/index.html', '树形 + 版本号 + 动态隐藏'],
    ['M', 'test/test_api_html.php', 'env baseUrl + 新接口 + 报告优化'],
    ['M', 'test/test_api_html_simp.php', 'env baseUrl + 新接口 + 报告优化'],
])
doc.add_paragraph()
p('-- 文档结束 --')

os.makedirs('docs', exist_ok=True)
out = 'docs/Devops-Glue-v2.2.0-版本升级说明.docx'
doc.save(out)
print(f'Done: {out}')
